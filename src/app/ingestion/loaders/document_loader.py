from __future__ import annotations

import asyncio
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.contracts.ingestion import BaseLoader, DocumentBlob, ParsedDocument


def _parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _parse_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


class DocumentLoader(BaseLoader):
    name = "document"
    _extensions = {".pdf", ".docx"}

    def supports(self, filename: str, content_type: str) -> bool:
        ext = Path(filename).suffix.lower()
        return ext in self._extensions or content_type in {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

    async def load(self, blob: DocumentBlob) -> ParsedDocument:
        ext = Path(blob.filename).suffix.lower()
        if ext == ".pdf":
            text = await asyncio.to_thread(_parse_pdf, blob.content_bytes)
        elif ext == ".docx":
            text = await asyncio.to_thread(_parse_docx, blob.content_bytes)
        else:
            text = blob.content_bytes.decode("utf-8", errors="ignore")

        return ParsedDocument(
            document_id=blob.document_id,
            text=text,
            metadata={
                **blob.metadata,
                "filename": blob.filename,
                "loader": self.name,
            },
        )
